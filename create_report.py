from docx import Document
from datetime import datetime
from overall_state import OverallState
import io

def create_document(state: OverallState) -> dict:
    """
    Node for creating the final Word document and returning it as a byte stream.
    This avoids saving the file to disk on the server.
    """
    ui_container = state.get("ui_container")
    summarized_dict = state["summarized_articles"]
    query = state["messages"][-1].content
    ui_container.write("- Creating analysis report...")

    doc = Document()
    doc.add_heading(f'Research Analysis: {query}', 0)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("")

    for publication, summary_dict in summarized_dict.items():
        if summary_dict:
            doc.add_heading(publication, level=1)
            doc.add_paragraph("")
            for title, summarized_article_dict in summary_dict.items():
                summary = summarized_article_dict.get("summary", "Summary not available.")
                doc.add_heading(title, level=2)
                doc.add_paragraph(summary)
                doc.add_paragraph(f"Source: {summarized_article_dict.get('link', 'Link not available.')}")
                doc.add_paragraph("")

    safe_query = "".join(c for c in query if c.isalnum() or c in (' ', '_')).rstrip()
    safe_query = safe_query.replace(' ', '_')
    file_name = f"research_analysis_{safe_query}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    return {
        "document_bytes": doc_io.getvalue(),
        "document_name": file_name
    }

